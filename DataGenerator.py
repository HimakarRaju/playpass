import pandas as pd
import random
from faker import Faker
import numpy as np
from tqdm import tqdm  # For progress tracking

# Setup
fake = Faker()
Faker.seed(42)
random.seed(42)
np.random.seed(42)

# Constants
num_rows = 50000
batch_size = 10000  # Process in batches to reduce memory usage

categories = {
    "Games": ["Puzzle", "RPG", "Action", "Casual", "Strategy"],
    "Productivity": ["Note Taking", "Calendar", "To-Do List"],
    "Social": ["Messaging", "Photo Sharing", "Live Streaming"],
    "Tools": ["File Manager", "Battery Saver", "Cleaner"],
    "Education": ["Language Learning", "Online Courses", "Kids Learning"],
}

transaction_types = ["Install", "In-App Purchase", "Subscription", "Ad Click", "Refund"]
genders = ["Male", "Female", "Other"]
income_levels = ["Low", "Medium", "High"]
device_types = [
    "Samsung Galaxy S23",
    "Pixel 7",
    "OnePlus 11",
    "Xiaomi 13",
    "Motorola Edge",
]
android_versions = ["11", "12", "13", "14"]
regions = ["North America", "Europe", "Southeast Asia", "South America", "Africa"]
play_pass_types = ["Free", "Basic", "Premium"]
days_of_week = [
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
]
seasons = ["Spring", "Summer", "Autumn", "Winter"]
countries = ["USA", "India", "UK", "Germany", "Brazil", "Japan", "Australia", "Canada", "France", "Mexico", "South Africa", "China"]
states = {
    "USA": ["California", "Texas", "New York", "Florida", "Illinois"],
    "India": ["Maharashtra", "Karnataka", "Delhi", "Tamil Nadu", "Uttar Pradesh"],
    "UK": ["England", "Scotland", "Wales", "Northern Ireland"],
    "Germany": ["Bavaria", "Berlin", "Hamburg", "North Rhine-Westphalia"],
    "Brazil": ["São Paulo", "Rio de Janeiro", "Minas Gerais", "Bahia"],
    "Japan": ["Tokyo", "Osaka", "Kyoto", "Hokkaido"],
    "Australia": ["New South Wales", "Victoria", "Queensland", "Western Australia"],
    "Canada": ["Ontario", "Quebec", "British Columbia", "Alberta"],
    "France": ["Île-de-France", "Provence-Alpes-Côte d'Azur", "Auvergne-Rhône-Alpes"],
    "Mexico": ["Mexico City", "Jalisco", "Nuevo León", "Puebla"],
    "South Africa": ["Gauteng", "Western Cape", "KwaZulu-Natal"],
    "China": ["Beijing", "Shanghai", "Guangdong", "Sichuan"]
}
app_tags = [
    "Offline",
    "Multiplayer",
    "Editor’s Choice",
    "Family Friendly",
    "AR Support",
]
age_ratings = ["Everyone", "Teen", "18+"]
languages = ["en-US", "hi-IN", "pt-BR", "de-DE", "en-GB"]


def generate_amount_spent(transaction_type):
    if transaction_type == "In-App Purchase":
        return round(random.uniform(0.99, 200), 2)
    elif transaction_type == "Subscription":
        return round(random.uniform(1.99, 50), 2)
    else:
        return 0.0


# Optimized data generation
def generate_batch(start_idx, end_idx):
    batch_data = []
    for i in range(start_idx, end_idx):
        country = random.choice(countries)
        state = random.choice(states[country])
        category = random.choice(list(categories.keys()))
        sub_category = random.choice(categories[category])
        app_name = f"{fake.catch_phrase().title()} {sub_category}"
        app_size = round(random.uniform(5, 500), 2)
        total_installs = random.randint(1000, 10000000)
        free_paid = random.choice(["Free", "Paid"])
        discount_applied = random.choice(["Yes", "No"])
        promo_code_used = random.choice(["Yes", "No"])
        base_price = round(random.uniform(0.99, 49.99), 2) if free_paid == "Paid" else 0.0
        final_price = (
            round(base_price * random.uniform(0.3, 0.8), 2)
            if free_paid == "Paid"
            and (discount_applied == "Yes" or promo_code_used == "Yes")
            else base_price
        )
        transaction_type = random.choice(transaction_types)

        row = {
            "ID": i + 1,
            "Name": fake.name(),
            "Email": fake.free_email() if country in ["USA", "UK", "Canada", "Australia"] 
                    else fake.email() if country == "India" 
                    else fake.ascii_free_email() if country == "Japan" 
                    else fake.ascii_email() if country == "China" 
                    else fake.company_email(),
            "Phone": fake.phone_number(),
            "Age": random.randint(13, 65),
            "Gender": random.choice(genders),
            "Income Level": random.choice(income_levels),
            "Device Type": random.choice(device_types),
            "Android Version": random.choice(android_versions),
            "App Name": app_name,
            "Developer": fake.company(),
            "Category": category,
            "Sub_Category": sub_category,
            "Free/Paid": free_paid,
            "In-App Purchases": random.choice(["Yes", "No"]),
            "App Size (MB)": app_size,
            "Total Installs": total_installs,
            "Transaction ID": fake.uuid4(),
            "Transaction Type": transaction_type,
            "App/Game Price": base_price,
            "Discount Applied": discount_applied,
            "Promo Code Used": promo_code_used,
            "Price Paid (with Coupon)": final_price,
            "Amount Spent on In-App Purchases": generate_amount_spent(transaction_type),
            "Time Spent (min)": random.randint(0, 300),
            "Session Count": random.randint(1, 50),
            "Time Since Last Use (days)": random.randint(0, 60),
            "Favorite Flag": random.choice(["Yes", "No"]),
            "Uninstalled": random.choice(["Yes", "No"]),
            "Date": fake.date_this_year(),
            "Time": fake.time(),
            "Day of Week": random.choice(days_of_week),
            "Weekend": random.choice(["Yes", "No"]),
            "Season": random.choice(seasons),
            "Rating": random.randint(1, 5),
            "Review Text": fake.sentence(nb_words=10),
            "Review Sentiment": random.choice(["Positive", "Neutral", "Negative"]),
            "Review Length": random.randint(20, 300),
            "Demographic Location": fake.city(),
            "State": state,
            "Country": country,
            "Region": random.choice(regions),
            "Play Pass User": random.choice(play_pass_types),
            "Subscription Duration": random.randint(0, 24),
            "Auto-Renew": random.choice(["Yes", "No"]),
            "App Tags": ", ".join(random.sample(app_tags, 2)),
            "Age Rating": random.choice(age_ratings),
            "Device Locale/Language": random.choice(languages),
        }
        batch_data.append(row)
    return batch_data

# Process data in batches
df = pd.DataFrame()
for batch_start in tqdm(range(0, num_rows, batch_size), desc="Generating data"):
    batch_end = min(batch_start + batch_size, num_rows)
    batch_data = generate_batch(batch_start, batch_end)
    batch_df = pd.DataFrame(batch_data)
    
    # Optimize memory usage by specifying dtypes
    numeric_cols = ['Age', 'Total Installs', 'Time Spent (min)', 'Session Count', 
                   'Time Since Last Use (days)', 'Rating', 'Review Length', 'Subscription Duration']
    batch_df[numeric_cols] = batch_df[numeric_cols].astype('int32')
    
    float_cols = ['App Size (MB)', 'App/Game Price', 'Price Paid (with Coupon)', 
                 'Amount Spent on In-App Purchases']
    batch_df[float_cols] = batch_df[float_cols].astype('float32')
    
    df = pd.concat([df, batch_df], ignore_index=True)

# Save to Excel with optimized engine
df.to_excel("full_playstore_dataset_50000.xlsx", index=False, engine='openpyxl')

# Optimized summary reporting
def generate_summary(df):
    summary = {
        'Total rows': len(df),
        'Total columns': len(df.columns),
        'Dataset shape': df.shape,
        'Memory usage (MB)': df.memory_usage(deep=True).sum() / (1024 * 1024),
        'Null values': df.isnull().sum().to_dict(),
        'Unique values': df.nunique().to_dict()
    }
    return summary

print("✅ Dataset generated and saved as full_playstore_dataset_50000.xlsx")
print("\n=== Optimized Summary ===")
summary = generate_summary(df)
for key, value in summary.items():
    print(f"{key}: {value}")

# Generate optimized report
from docx import Document

def create_report(df, filename="output.docx"):
    doc = Document()
    doc.add_heading('Play Store Dataset Summary Report', level=1)
    
    # Add optimized summary
    summary = generate_summary(df)
    doc.add_heading('Key Metrics', level=2)
    for key, value in summary.items():
        if isinstance(value, dict):
            doc.add_paragraph(f"{key}:")
            for k, v in value.items():
                doc.add_paragraph(f"  {k}: {v}", style='ListBullet')
        else:
            doc.add_paragraph(f"{key}: {value}")
    
    # Add sample data
    doc.add_heading('Sample Data', level=2)
    doc.add_paragraph(str(df.sample(5)))
    
    doc.save(filename)
    print(f"✅ Report generated and saved as {filename}")

create_report(df)

# Optimized visualization
def visualize_data(df):
    import matplotlib.pyplot as plt
    import seaborn as sns
    
    # Set style
    sns.set_style("whitegrid")
    
    # Create subplots
    fig, axes = plt.subplots(2, 2, figsize=(15, 10))
    
    # Plot 1: App size distribution
    sns.histplot(df['App Size (MB)'], bins=30, kde=True, ax=axes[0,0])
    axes[0,0].set_title('App Size Distribution')
    
    # Plot 2: Category distribution
    sns.countplot(y='Category', data=df, ax=axes[0,1], 
                 order=df['Category'].value_counts().index)
    axes[0,1].set_title('App Categories')
    
    # Plot 3: Rating distribution
    sns.boxplot(x='Category', y='Rating', data=df, ax=axes[1,0])
    axes[1,0].set_title('Ratings by Category')
    axes[1,0].tick_params(axis='x', rotation=45)
    
    # Plot 4: Time spent vs rating
    sns.scatterplot(x='Time Spent (min)', y='Rating', data=df, ax=axes[1,1], alpha=0.3)
    axes[1,1].set_title('Time Spent vs Rating')
    
    plt.tight_layout()
    plt.show()

visualize_data(df)
